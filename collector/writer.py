# -*- coding: utf-8 -*-
from __future__ import print_function

import os
import random
import locale

from datetime import datetime, timedelta
from sqlalchemy.sql.expression import func, and_
from docx import Document
from docx.shared import Inches, Pt
from dateutil import parser as date_parser

from models import DBSession, Apartment, Link

locale.setlocale(locale.LC_ALL, 'sv_SE.utf8')

def write_to_one_file(
    coupon,
    coupon_expiry,
    customer_name,
    customer_formal_address,
    coupon_value=500,
    min_order_value=2500,
    doc_template='LetterTemplate.docx',
    file_save_name=None,
    save_path='.'
):
    """
    Creates a marketing letter (docx) for a customer with given data
    Parameters
    ----------
    coupon : str
    coupon_expiry : str
    customer_name : str
    customer_formal_address : str
    coupon_value : int
    min_order_value : int
    doc_template : str
    file_save_name : str|None
    save_path : str

    """

    tell = '070-322 55 55'
    if len(customer_name) > 25:
        name_split = customer_name.split(' ')
        customer_name = name_split[0] + ' ' + name_split[-1]
    customer_name = customer_name
    address_split = customer_formal_address.split(',')
    customer_address = address_split[0].strip()  # u'Karlavägen 68'
    customer_zip_city = address_split[1].strip()  # u'11 459 Stockholm'
    if file_save_name is None:
        file_save_name = coupon + '.docx'

    document = Document(doc_template)
    tbl1 = document.tables[0]
    col1 = tbl1.columns[0]

    cpn_box_val_run = col1.cells[0].paragraphs[0].add_run(str(coupon_value) + ' kr')
    cpn_box_val_font = cpn_box_val_run.font
    cpn_box_val_font.size = Pt(54)
    coupon_val = col1.cells[2].paragraphs[0].add_run(coupon)
    coupon_val_font = coupon_val.font
    coupon_val_font.size = Pt(18)
    col2 = tbl1.columns[1]
    col2_cell2 = col2.cells[0]
    address_cell = col2_cell2.tables[0].rows[0].cells[1]
    address_cell.text = u'{name}\n{address}\n{zipcity}'.format(
        name=customer_name,
        address=unicode(customer_address),
        zipcity=customer_zip_city
    )
    address_cell_font = address_cell.paragraphs[0].runs[0].font
    address_cell_font.size = Pt(12)

    pars = document.paragraphs
    for par in pars:
        par.paragraph_format.right_indent = Inches(-1.3)

    p9 = pars[9]
    p9_appended = p9.add_run(
        u'Vi har bifogat en kupong med värde {value} kr som en välkomstgåva för din bokning med cleanJoy.se.'.format(
            value=coupon_value
        ),
    )

    p19 = pars[19]
    p19_appended = p19.add_run(
        u'Boka en flyttstädning idag och få {value} kr rabatt på beställningar över {max_value} kr. Erbjudandet gäller till den {expiry}.'
        .format(
            value=coupon_value,
            expiry=coupon_expiry,
            max_value=min_order_value
        )
    )  # bold
    p19_appended_font = p19_appended.font
    p19_appended_font.bold = True

    p20 = pars[20]
    p20_appended = p20.add_run(
        u'Ring vår kundtjänst på {tell} så hjälper de dig med frågor angående bokning och tjänster. Eller gå direct till '.format(
            tell=tell
        )
    )
    p20_appended = p20.add_run(
        u'cleanjoy.se/aktivera'
    )
    p20_appended_font = p20_appended.font
    p20_appended_font.underline = True
    p18_appended = p20.add_run(
        u' för att boka din städning.'
    )
    p27 = pars[27]
    p27_appended = p27.add_run(
        u'Erbjudandet gäller t.o.m {expiry}.'.format(
            expiry=date_parser.parse(coupon_expiry).date().strftime('%d %B %Y')
        )
    )
    save_path_name = os.path.join(
        save_path,
        file_save_name
    )
    document.save(save_path_name)


def generate_coupon_code():
        code_chars = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'

        def four_chars():
            four_char = ''
            for i in range(4):
                four_char += random.choice(code_chars)
            return four_char
        coupon_code = 'LTTR-' + '-'.join([four_chars() for i in range(3)])
        return coupon_code


def log_coupons(
    save_path,
    data,
    log_name='coupons.txt'
):
    """
    Writes newly generated coupons to a file along with other info, e.g expiry, etc.

    Parameters
    ----------
    save_path : str
    log_name : str
    data : str
    """

    log_file = os.path.join(
        save_path,
        log_name,
    )
    with open(log_file, mode='a') as log:
        log.write(data + '\n')


def write_to_file(
    save_path='/media/spartonia/DATA/Company_stuff/Marketing/Letters',
    published_days_ago=30,
    date_info_collected=str((datetime.now().date() - timedelta(days=-1))),  # yesterday
    price_range=(0, 20000000),
    living_area_range=(0, 500),
    merge_to_pdf=True
):
    # TODO: add county selection option
    """
    Writes marketing letters.

    Parameters
    ----------
    save_path: str
    published_days_ago:  int
    date_info_collected: str 'YYYY-mm-dd'
    price_range: (int, int)
    living_area_range: (int, int)
    merge_to_pdf : bool

    """
    reporting_day_save_path = os.path.join(
        save_path,
        str(datetime.now().date())
    )
    if not os.path.exists(reporting_day_save_path):
        os.makedirs(reporting_day_save_path)

    save_path_reporting_day_collected_day = os.path.join(
        reporting_day_save_path,
        date_info_collected
    )
    if not os.path.exists(save_path_reporting_day_collected_day):
        os.mkdir(save_path_reporting_day_collected_day)

    date_collected = date_parser.parse(date_info_collected).date()
    session = DBSession()
    apts = session.query(Apartment).join(Link).filter(
        Apartment.processed.isnot(True),
        Apartment.owner.isnot(None),
        and_(Apartment.price >= price_range[0], Apartment.price <= price_range[1]),
        and_(Apartment.living_area >= living_area_range[0], Apartment.living_area <= living_area_range[1]),
        Apartment.price.isnot(None),
        Apartment.postal_code.isnot(None),
        Link.timestamp.between(
            str(date_collected), str(date_collected + timedelta(days=1))),
        Link.date.between(str(datetime.now().date() + timedelta(days=-10)), str(datetime.now().date())),
    )
    file_count = 0
    for i, apt in enumerate(apts):
        coupon_code = generate_coupon_code()
        expiry = str(datetime.now().date() + timedelta(days=120))
        coupon_value = 500
        print(apt.owner)
        try:
            write_to_one_file(
                coupon=coupon_code,
                coupon_expiry=expiry,
                coupon_value=coupon_value,
                customer_name=apt.owner,
                customer_formal_address=apt.formal_address,
                save_path=save_path_reporting_day_collected_day
            )
            log_data = coupon_code + ';' + expiry + ';' + str(coupon_value)
            log_coupons(
                reporting_day_save_path,
                data=log_data
            )
            file_count += 1
            apt.processed = True
            try:
                session.commit()
            except Exception as e:
                print(e.message)
                session = DBSession()
        except Exception as e:
            print('Error creating letter for', apt.formal_address, ',', apt.owner, coupon_code, expiry)
            print('Error message:', e.message)
    print('Done!')
    print(file_count, 'letters created.')


if __name__ == '__main__':
    # write_to_file(date_info_collected='2016-02-04')
    pass

    # coupon_code = generate_coupon_code()
    # print(coupon_code)